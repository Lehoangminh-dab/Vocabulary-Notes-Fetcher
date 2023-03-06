def docx_test():
    import docx
    from docx.shared import RGBColor
    from docx.shared import Inches,Pt
    document=docx.Document('demo.docx')
    def test(document):
        table=document.add_table(rows=1,cols=3,style='Table Grid')
        content='Bowser eating shit'
        def write_content(row_num,cell_num,content):
            content_paragraph=table.rows[row_num-1].cells[cell_num-1].add_paragraph().add_run(content)
            content_paragraph.font.name='Merriweather'
            content_paragraph.font.size=Pt(12)
            content_paragraph.font.color.rgb=RGBColor(255,0,0)
        write_content(1,3,content)
        document.save('demo.docx')
    test(document)
    font=document.tables[1].rows[0].cells[2].add_paragraph().add_run("Bowser's son eating shit")
    font.italic=True
    document.save('demo.docx')


def file_read_test():
    import sys
    import os
    sys.path.append(os.path.abspath('D:\Minh\Code\Translate god\Version 1.0\V1 Beta\Main\Dependencies'))
    file=open('D:\Minh\Code\Translate god\Version 1.0\V1 Beta\Main\Dependencies\Fonts.txt','r')
    print(file.readlines())


def input_conversion_test():
    try:
        user_input=int(input())
    except ValueError:
        print("Value invalid, please type an integer")
        user_input=input_conversion_test()
    return user_input
def main():
    print(input_conversion_test())

main()