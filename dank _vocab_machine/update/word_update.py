from docx.api import Document
from docx.shared import RGBColor
import docx
from docx.shared import Inches,Pt
#The selected entries will have the following format:
#{Word_Type1:[{'endef':,'vndef':,examples:[]},{...}],Word_Type2:}

def update(document_name,word_id,word_data):
    document = docx.Document(document_name)
    fonts = open('D:\Programming Projects\Full Fledged Projects\Apps\Translate god\dank _vocab_machine\dependencies\Fonts.txt','r').read()
    font_file = fonts.split('\n')
#CONSTANTS
    #TABLE WIDTH
    _1ST_COLUMN_WIDTH = Inches(0.521)
    _2ND_COLUMN_WIDTH = Inches(1.865)
    _3RD_COLUMN_WIDTH = Inches(4.063)
    #title ROW FONTS
    TITLE_NO_FONT = font_file[2:9]
    TITLE_WORD_FONT = font_file[9:16]
    TITLE_DEFINITION_FONT = font_file[16:23]
    #1ST COLUMN FONTS
    WORD_COUNT_FONT = font_file[24:31]
    #2ND COLUMN FONTS
    WORD_ID_FONT = font_file[32:39]
    WORD_PRONUNCIATION_FONT = font_file[39:46]
    WORD_TYPES_FONT = font_file[46:53]
    #3RD COLUMN FONTS
    ENTRY_WORD_TYPE_FONT = font_file[54:61]
    ENTRY_EN_DEFINITION_FONT = font_file[61:68]
    ENTRY_EG_FONT = font_file[68:75]
    ENTRY_EXAMPLES_FONT = font_file[75:82]
    ENTRY_VN_DEFINITION_FONT = font_file[82:88]

    #MAIN
    def update_main():
        if len(document.tables) == 0:
            table_id = create_table()
        else:
            table_id = document.tables[0]
        write_word_data(table_id)
        set_col_widths(table_id)
        document.save(document_name)


    #Some functions for writing text
    def apply_font(text,font_data):
        def apply_font_main():   
            text_font_values = get_font_values()
            text.font.name = text_font_values[0]
            text.font.size = Pt(int(text_font_values[1]))
            color_value = text_font_values[2]
            text.font.color.rgb = RGBColor(int(color_value[0]),int(color_value[1]),int(color_value[2]))
            text.font.bold = text_font_values[3] == 'True'
            text.font.italic = text_font_values[4] == 'True'
            text.font.underline = text_font_values[5] == 'True'  


        def get_font_values():
            font_values = []
            for font_value in font_data:
                font_values.append(font_value.split(':')[1])
    
            def get_color_values(font_color_values):
                color_value = font_color_values.split(',')
                return color_value

            font_values[2] = get_color_values(font_values[2])
            return font_values

        apply_font_main()


    def write_paragraph(row_id,cell_id,content,content_font):
        text = row_id.cells[cell_id].add_paragraph().add_run(content)
        apply_font(text,content_font)              


    def write_list(row_id,cell_id,list_content,list_font,list_style):
        list_text = row_id.cells[cell_id].add_paragraph(style = list_style).add_run(list_content)
        apply_font(list_text,list_font)
    #End

    #Start of Logic
    def create_table():
        table = document.add_table(rows = 1,cols = 3,style = 'Table Grid')
        table.autofit = False
        titles_row = table.rows[0]
        #title:"No"
        title_no_content = "No"
        title_no_font = TITLE_NO_FONT
        title_no = write_paragraph(titles_row,0,title_no_content,title_no_font)
        #title:"Word"
        title_word_content = "Word"
        title_word_font = TITLE_WORD_FONT
        title_word = write_paragraph(titles_row,1,title_word_content,title_word_font)
        #title:"Definition"
        title_definition_content = "Definition"
        title_definition_font = TITLE_DEFINITION_FONT
        title_definition = write_paragraph(titles_row,2,title_definition_content,title_definition_font)
        return table

    def write_word_data(table_id):
        def write_word_data_main():
            if len(word_data) == 0:
                return 
            else:       
                write_word_count()
                write_word()
                write_word_definition()
        
        new_row = table_id.add_row()

        def add_row_paragraph(cell_id,content,content_font):
            write_paragraph(new_row,cell_id,content,content_font)

        def write_word_count():
            content = str(len(table_id.rows)-1)
            content_font = WORD_COUNT_FONT
            add_row_paragraph(0,content,content_font)


        def write_word():#Write the 2nd column entry
            def write_word_main():
                write_word_id()
                write_word_pronunciation()
                write_word_types()


            def write_word_id():
                content = '"'+word_id.capitalize()+'"'
                content_font = WORD_ID_FONT
                add_row_paragraph(1,content,content_font)

            def write_word_pronunciation():
                for pronunciation in word_data['pronunciations']:
                    content = '/'+pronunciation+'/'
                    content_font = WORD_PRONUNCIATION_FONT
                    add_row_paragraph(1,content,content_font)

            def write_word_types():
                content = '('
                for word_type in list(word_data):
                    if word_type == 'pronunciations':
                        break
                    else:
                        content+= word_type+','
                        pos = len(content)-1
                content = content[:pos] + content[(pos+1):]#delete the excess punctuation
                content+= ')'
                content_font = WORD_TYPES_FONT
                add_row_paragraph(1,content,content_font)

            write_word_main()


        def write_word_definition():
            def write_word_definition_main():
                for word_type in list(word_data):
                    if word_type == 'pronunciations':
                        break
                    else:
                        write_definitions_entry(word_type)


            def write_definitions_entry(word_type):
                def write_definitions_entry_main():
                    write_entry_word_type()
                    for entry_content in word_data[word_type]:
                        write_entry_content(entry_content)


                def write_entry_word_type():
                    content = '*'+word_type+':'
                    content_font = ENTRY_WORD_TYPE_FONT
                    add_row_paragraph(2,content,content_font)

                def write_entry_content(entry_content):
                    def write_entry_content_main():
                        write_definition()
                        write_examples()

                    def write_definition():
                        def write_definition_main():
                            write_en_definition()
                            write_vn_definition()


                        definition_count = word_data[word_type].index(entry_content)+1
                        def write_en_definition():
                            content = str(definition_count) + ')' + entry_content['endef'] + '.'
                            content_font = ENTRY_EN_DEFINITION_FONT
                            add_row_paragraph(2,content,content_font)
                        
                        def write_vn_definition():
                            #Format:"(VN:...)"
                            vn_definitions = entry_content['vndef']
                            content = "(VN: "
                            for vn_definition in vn_definitions:
                                if vn_definitions.index(vn_definition) == len(vn_definitions) - 1:
                                    content += vn_definition
                                else:
                                    content += vn_definition + ", "
                            content += ")"

                            content_font = ENTRY_VN_DEFINITION_FONT
                            add_row_paragraph(2, content, content_font)
                            

                        write_definition_main()



                    def write_examples():
                        def write_examples_main():
                            write_title()
                            write_examples_content()
                            
                        def write_title():
                            content = '*E.g:'
                            content_font = ENTRY_EG_FONT
                            add_row_paragraph(2,content,content_font)

                        def write_examples_content():
                            for example in entry_content['examples']:
                                content = example+'.'
                                content_font = ENTRY_EXAMPLES_FONT
                                write_list(new_row,2,content,content_font,'List Bullet')
                        
                        write_examples_main()

                    write_entry_content_main()

                write_definitions_entry_main()

            write_word_definition_main()

        write_word_data_main()
    
    def set_col_widths(table):
        widths  =  (_1ST_COLUMN_WIDTH,_2ND_COLUMN_WIDTH,_3RD_COLUMN_WIDTH)
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width  =  width


    update_main()
                   
                    






    